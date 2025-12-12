#!/usr/bin/env python3
"""
Générateur de templates CV professionnels AVEC DESIGN
Version améliorée avec couleurs, formatage et styles
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_background(cell, fill):
    """Ajoute une couleur de fond à une cellule de tableau"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_colored_heading(doc, text, level, color_rgb):
    """Ajoute un titre avec couleur"""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
    return heading

def create_classic_template_1():
    """CV Traditionnel Suisse - Bleu classique"""
    doc = Document()
    
    # Couleur bleu classique
    BLUE = (13, 71, 161)  # Bleu foncé professionnel
    
    # En-tête avec fond bleu
    header_table = doc.add_table(rows=1, cols=1)
    header_table.autofit = False
    header_table.allow_autofit = False
    header_cell = header_table.rows[0].cells[0]
    set_cell_background(header_cell, '0D47A1')
    
    # Nom en blanc dans l'en-tête
    name_para = header_cell.paragraphs[0]
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run('PRÉNOM NOM')
    name_run.font.size = Pt(28)
    name_run.font.color.rgb = RGBColor(255, 255, 255)
    name_run.bold = True
    
    # Contact
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.add_run('Adresse • Ville, Code Postal\nTéléphone: +41 XX XXX XX XX • Email: prenom.nom@email.ch\nLinkedIn: linkedin.com/in/votre-profil')
    contact_run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Profil avec titre bleu
    add_colored_heading(doc, 'PROFIL PROFESSIONNEL', 1, BLUE)
    doc.add_paragraph(
        'Professionnel expérimenté avec [X] années d\'expérience dans [domaine]. '
        'Spécialisé en [compétences clés]. Recherche un poste de [titre du poste] '
        'pour contribuer à [objectif professionnel].'
    )
    
    # Expérience
    add_colored_heading(doc, 'EXPÉRIENCE PROFESSIONNELLE', 1, BLUE)
    
    exp1 = doc.add_paragraph()
    exp1_title = exp1.add_run('Titre du Poste\n')
    exp1_title.bold = True
    exp1_title.font.size = Pt(12)
    exp1_company = exp1.add_run('Nom de l\'Entreprise, Ville • MM/AAAA - MM/AAAA\n')
    exp1_company.italic = True
    exp1_company.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph('• Responsabilité ou réalisation majeure', style='List Bullet')
    doc.add_paragraph('• Projet important avec résultats mesurables', style='List Bullet')
    doc.add_paragraph('• Contribution à l\'équipe ou l\'entreprise', style='List Bullet')
    
    doc.add_paragraph()
    
    # Formation
    add_colored_heading(doc, 'FORMATION', 1, BLUE)
    edu = doc.add_paragraph()
    edu_title = edu.add_run('Diplôme / Titre de Formation\n')
    edu_title.bold = True
    edu_inst = edu.add_run('Nom de l\'Institution, Ville • Année d\'obtention')
    edu_inst.italic = True
    
    # Compétences
    add_colored_heading(doc, 'COMPÉTENCES', 1, BLUE)
    skills_table = doc.add_table(rows=3, cols=2)
    skills_table.style = 'Light Grid Accent 1'
    
    # En-têtes du tableau
    header_cells = skills_table.rows[0].cells
    header_cells[0].text = 'Catégorie'
    header_cells[1].text = 'Compétences'
    for cell in header_cells:
        set_cell_background(cell, 'E3F2FD')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    skills_table.rows[1].cells[0].text = 'Techniques'
    skills_table.rows[1].cells[1].text = 'Compétence 1, Compétence 2, Compétence 3'
    skills_table.rows[2].cells[0].text = 'Langues'
    skills_table.rows[2].cells[1].text = 'Français (Natif), Anglais (C1), Allemand (B2)'
    
    return doc

def create_modern_template_1():
    """CV Tech Moderne - Bleu électrique"""
    doc = Document()
    
    ELECTRIC_BLUE = (33, 150, 243)
    DARK_GRAY = (66, 66, 66)
    
    # Header moderne
    header = doc.add_paragraph()
    header_run = header.add_run('< PRÉNOM NOM />')
    header_run.font.size = Pt(26)
    header_run.font.color.rgb = RGBColor(*ELECTRIC_BLUE)
    header_run.bold = True
    
    subtitle = doc.add_paragraph('Développeur Full Stack | Tech Enthusiast')
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(*DARK_GRAY)
    
    # Ligne de séparation
    separator = doc.add_paragraph('─' * 70)
    separator.runs[0].font.color.rgb = RGBColor(*ELECTRIC_BLUE)
    
    # Contact
    contact = doc.add_paragraph('💻 github.com/username • 🌐 portfolio.com • 📧 dev@email.ch • 📱 +41 XX XXX XX XX')
    contact.runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    
    # Sections avec titres colorés
    add_colored_heading(doc, '// À PROPOS', 1, ELECTRIC_BLUE)
    doc.add_paragraph(
        'Développeur passionné avec [X] ans d\'expérience en développement web et mobile. '
        'Spécialisé en [technologies]. Contributeur open-source.'
    )
    
    add_colored_heading(doc, '// STACK TECHNIQUE', 1, ELECTRIC_BLUE)
    
    # Tableau de compétences avec fond coloré
    stack_table = doc.add_table(rows=4, cols=2)
    stack_table.style = 'Light Grid'
    
    categories = [
        ('Frontend', 'React, Vue.js, TypeScript, HTML5, CSS3'),
        ('Backend', 'Node.js, Python, Django, Express'),
        ('Database', 'PostgreSQL, MongoDB, Redis'),
        ('DevOps', 'Docker, Kubernetes, CI/CD, AWS')
    ]
    
    for i, (cat, skills) in enumerate(categories):
        row = stack_table.rows[i]
        cat_cell = row.cells[0]
        skills_cell = row.cells[1]
        
        set_cell_background(cat_cell, 'E1F5FE')
        cat_cell.text = cat
        cat_cell.paragraphs[0].runs[0].font.bold = True
        
        skills_cell.text = skills
    
    add_colored_heading(doc, '// EXPÉRIENCE', 1, ELECTRIC_BLUE)
    
    exp = doc.add_paragraph()
    exp_title = exp.add_run('Senior Developer @ TechCorp\n')
    exp_title.bold = True
    exp_title.font.size = Pt(12)
    exp_date = exp.add_run('Genève, CH • 2022 - Présent\n')
    exp_date.italic = True
    exp_date.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph('→ Développement d\'applications web scalables', style='List Bullet')
    doc.add_paragraph('→ Architecture microservices et APIs', style='List Bullet')
    
    return doc

def create_creative_template_1():
    """CV Design Créatif - Violet/Rose"""
    doc = Document()
    
    PURPLE = (156, 39, 176)
    PINK = (233, 30, 99)
    
    # Header créatif avec fond dégradé simulé
    header_table = doc.add_table(rows=1, cols=1)
    header_cell = header_table.rows[0].cells[0]
    set_cell_background(header_cell, '9C27B0')
    
    name_para = header_cell.paragraphs[0]
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run('✦ PRÉNOM NOM ✦')
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = RGBColor(255, 255, 255)
    name_run.bold = True
    
    # Sous-titre
    subtitle = doc.add_paragraph('Designer Graphique | Créatif Visuel')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(*PURPLE)
    
    # Contact
    contact = doc.add_paragraph('🎨 portfolio.com  •  📧 design@email.ch  •  📱 +41 XX XXX XX XX')
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.runs[0].font.size = Pt(9)
    
    separator = doc.add_paragraph('✦ ✦ ✦ ✦ ✦ ✦ ✦ ✦')
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator.runs[0].font.color.rgb = RGBColor(*PINK)
    
    doc.add_paragraph()
    
    # Sections
    add_colored_heading(doc, '✦ À PROPOS', 1, PURPLE)
    doc.add_paragraph(
        'Designer créatif avec [X] ans d\'expérience en design graphique et branding. '
        'Passionné par la création d\'identités visuelles fortes.'
    )
    
    add_colored_heading(doc, '✦ EXPERTISE CRÉATIVE', 1, PURPLE)
    
    # Tableau coloré
    expertise_table = doc.add_table(rows=4, cols=1)
    expertise_items = [
        '🎨 Design Graphique: Identité visuelle, Logo, Print',
        '💻 Design Digital: UI/UX, Web design, Motion',
        '🛠️ Outils: Adobe Creative Suite, Figma, Sketch',
        '🎯 Spécialités: Branding, Typography, Illustration'
    ]
    
    for i, item in enumerate(expertise_items):
        cell = expertise_table.rows[i].cells[0]
        if i % 2 == 0:
            set_cell_background(cell, 'F3E5F5')
        cell.text = item
    
    add_colored_heading(doc, '✦ EXPÉRIENCE', 1, PURPLE)
    
    exp = doc.add_paragraph()
    exp_title = exp.add_run('Senior Designer @ Agence Créative\n')
    exp_title.bold = True
    exp_title.font.color.rgb = RGBColor(*PINK)
    exp_date = exp.add_run('Genève • 2021 - Présent')
    exp_date.italic = True
    
    doc.add_paragraph('→ Création d\'identités visuelles pour 20+ clients', style='List Bullet')
    doc.add_paragraph('→ Direction artistique de campagnes', style='List Bullet')
    
    return doc

# Créer templates simplifiés pour les autres
def create_simple_template(title, color_hex, color_rgb):
    """Template générique avec couleur personnalisée"""
    doc = Document()
    
    # Header coloré
    header_table = doc.add_table(rows=1, cols=1)
    header_cell = header_table.rows[0].cells[0]
    set_cell_background(header_cell, color_hex)
    
    name_para = header_cell.paragraphs[0]
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run('PRÉNOM NOM')
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = RGBColor(255, 255, 255)
    name_run.bold = True
    
    # Contact
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('Email: prenom.nom@email.ch | Tél: +41 XX XXX XX XX | Ville, Suisse')
    
    doc.add_paragraph()
    
    # Profil
    add_colored_heading(doc, 'PROFIL PROFESSIONNEL', 1, color_rgb)
    doc.add_paragraph('Professionnel avec [X] années d\'expérience dans [domaine]. Recherche [objectif].')
    
    # Expérience
    add_colored_heading(doc, 'EXPÉRIENCE', 1, color_rgb)
    exp = doc.add_paragraph()
    exp.add_run('Titre du Poste\n').bold = True
    exp.add_run('Entreprise • 2022 - Présent\n').italic = True
    doc.add_paragraph('• Responsabilité principale', style='List Bullet')
    doc.add_paragraph('• Réalisation mesurable', style='List Bullet')
    
    # Formation
    add_colored_heading(doc, 'FORMATION', 1, color_rgb)
    doc.add_paragraph('Diplôme en [Domaine] - Institution (Année)')
    
    # Compétences
    add_colored_heading(doc, 'COMPÉTENCES', 1, color_rgb)
    skills_table = doc.add_table(rows=2, cols=2)
    skills_table.style = 'Light Grid'
    
    header_cells = skills_table.rows[0].cells
    header_cells[0].text = 'Catégorie'
    header_cells[1].text = 'Détails'
    for cell in header_cells:
        set_cell_background(cell, color_hex + '40')  # Version plus claire
        cell.paragraphs[0].runs[0].font.bold = True
    
    skills_table.rows[1].cells[0].text = 'Techniques'
    skills_table.rows[1].cells[1].text = 'Compétence 1, Compétence 2, Compétence 3'
    
    return doc

# Générer tous les templates
templates_config = [
    ('template-classique-1.docx', create_classic_template_1),
    ('template-classique-2.docx', lambda: create_simple_template('Corporate', '1565C0', (21, 101, 192))),
    ('template-classique-3.docx', lambda: create_simple_template('Business', '424242', (66, 66, 66))),
    ('template-moderne-1.docx', create_modern_template_1),
    ('template-moderne-2.docx', lambda: create_simple_template('Startup', 'FF6F00', (255, 111, 0))),
    ('template-moderne-3.docx', lambda: create_simple_template('Minimal', '37474F', (55, 71, 79))),
    ('template-creatif-1.docx', create_creative_template_1),
    ('template-creatif-2.docx', lambda: create_simple_template('Coloré', 'E91E63', (233, 30, 99))),
    ('template-creatif-3.docx', lambda: create_simple_template('Portfolio', '00897B', (0, 137, 123))),
]

os.makedirs('templates', exist_ok=True)

print("🎨 Génération des templates CV avec DESIGN...")
for filename, create_func in templates_config:
    doc = create_func()
    filepath = os.path.join('templates', filename)
    doc.save(filepath)
    print(f"✅ Créé: {filename}")

print("\n✨ Tous les templates avec design ont été créés!")
print("📁 Emplacement: ./templates/")
